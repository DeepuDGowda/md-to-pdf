# merge.py
from __future__ import annotations
import os, re
from io import BytesIO
from pathlib import Path
from typing import Dict, Optional, List, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Emu
from docx.table import _Cell, Table

# Markdown → DOCX bridge
from parser import md_file_to_docx

# ===================== STYLE MAP (match your template) =====================

STYLE_MAP: Dict[str, str] = {
    # Headings
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Heading 4": "Heading 4",

    # Lists
    "List Bullet":  "Bullet Point 1",  # we’ll also fallback if this doesn’t exist
    "List Bullet 2": "Bullet Point 2",
    "Bullet Point 1": "Bullet Point 1",
    "Bullet Point 2": "Bullet Point 2",
    "List Number":  "List Number1",
    "List Number 2":"List Number1",

    # Notes
    "Note": "Note",
    "Note_1": "Note_1",

    # Body fallbacks
    "Normal": "Body Text1",
    "Body Text": "Body Text1",
    "Default": "Body Text1",
}

# For numeric headings like "1.", "1.1", …
LEVEL_TAG_TO_STYLE = {"H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"}

# Preserve list indentation if found
BULLET_LEVEL_TO_STYLE = {0: "Bullet Point 1", 1: "Bullet Point 2"}

DEFAULT_BODY_STYLE = "Body Text1"
STEP_STYLE_NAME   = "Step 1"      # plain paragraphs should become "Step 1"
TABLE_STYLE_NAME  = "Table Grid"  # adjust if your template uses a different table style

# ===================== IMAGE / CAPTION POLICY =====================

CAPTION_STYLE_NAME = "Image Description"
CAPTION_PREFIX = "Figure"
NORMAL_WIDTH_CM = 15.0
SMALL_WIDTH_CM  = 9.0
SMALL_SOURCE_THRESHOLD_CM = 10.0

SUPPORTED_IMG_MIMES = {
    "image/png","image/jpeg","image/jpg","image/gif","image/bmp","image/tiff",
}

try:
    from PIL import Image  # optional
except Exception:
    Image = None

# ===================== SETTINGS / UTILITIES =====================

def _set_update_fields_on_open(doc: Document) -> None:
    settings_el = doc.settings._element
    update_el = settings_el.find(qn('w:updateFields'))
    if update_el is None:
        update_el = OxmlElement('w:updateFields')
        settings_el.append(update_el)
    update_el.set(qn('w:val'), 'true')

def _append_page_break(dst_doc: Document) -> None:
    try:
        dst_doc.add_page_break()
    except Exception:
        pass

def _get_ilvl(paragraph) -> Optional[int]:
    try:
        ilvl_nodes = paragraph._p.xpath('.//w:pPr/w:numPr/w:ilvl')
        if ilvl_nodes:
            return int(ilvl_nodes[0].get(qn('w:val')))
    except Exception:
        pass
    return None

def _looks_like_heading(text: str) -> Optional[str]:
    s = (text or "").strip()
    if s and s[0].isdigit():
        first = s.split()[0]
        dots = first.count(".")
        if dots == 0: return "H1"
        if dots == 1: return "H2"
        if dots >= 2: return "H3"
    return None

_STEP_MARKER = re.compile(r'^\s*(?:step\b|\d+[\.\)]|\(\d+\))\s+', re.IGNORECASE)

def _looks_like_procedure(raw_para) -> bool:
    if _get_ilvl(raw_para) is not None:
        return False
    text = (raw_para.text or "").strip()
    if not text:
        return False
    if _looks_like_heading(text):
        return False
    sty = (getattr(raw_para.style, "name", "") or "").lower()
    if "step" in sty:
        return True
    if _STEP_MARKER.match(text):
        return True
    return False

def _first_existing_style(doc: Document, names: List[str]) -> Optional[str]:
    for name in names:
        try:
            if name in doc.styles:  # python-docx implements membership check
                return name
        except Exception:
            pass
    return None

def _choose_style_for_paragraph(tpl: Document, raw_para) -> str:
    """
    Decide template style for this paragraph.
    Critical rule: if the source is a bullet/list, NEVER turn it into Step 1.
    """
    raw_name = (raw_para.style.name or "").strip()
    text     = (raw_para.text or "").strip()

    # --- 1) If it’s a bullet/numbered list, choose a bullet style with fallbacks
    if raw_name in ("List Bullet", "List Bullet 2", "Bullet Point 1", "Bullet Point 2") or _get_ilvl(raw_para) is not None:
        bullet_choice = _first_existing_style(
            tpl,
            [STYLE_MAP.get(raw_name, ""), "Bullet Point 1", "List Bullet", "List Paragraph"]
        )
        if bullet_choice:
            return bullet_choice

    # --- 2) Direct style mapping if present and exists
    if raw_name in STYLE_MAP and STYLE_MAP[raw_name] in tpl.styles:
        return STYLE_MAP[raw_name]

    # --- 3) Numeric heading detection like "2.", "2.1" …
    tag = _looks_like_heading(text)
    if tag and LEVEL_TAG_TO_STYLE.get(tag) in tpl.styles:
        return LEVEL_TAG_TO_STYLE[tag]

    # --- 4) Otherwise, plain paragraphs become Step 1 (procedure)
    if text and STEP_STYLE_NAME in tpl.styles:
        return STEP_STYLE_NAME

    # --- 5) Fallback to body
    return _first_existing_style(tpl, [DEFAULT_BODY_STYLE, raw_name]) or raw_name

# ===================== IMAGE HANDLING =====================

def _extract_inline_images_from_run(raw_doc: Document, run) -> List[Tuple[bytes, Optional[int], Optional[int], Optional[str]]]:
    results: List[Tuple[bytes, Optional[int], Optional[int], Optional[str]]] = []
    blip_elems = run._r.xpath('.//a:blip')
    if not blip_elems:
        return results

    wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'

    for blip in blip_elems:
        rId = blip.get(qn('r:embed'))
        if not rId:
            continue
        part = raw_doc.part.related_parts.get(rId)
        if part is None:
            continue
        content_type = getattr(part, 'content_type', None)
        if not content_type or not content_type.startswith("image/"):
            continue
        blob = getattr(part, 'blob', None)
        if not blob:
            continue

        cx = cy = None
        node = None
        inline = run._r.xpath('.//wp:inline')
        anchor = run._r.xpath('.//wp:anchor')
        if inline:
            node = inline[0]
        elif anchor:
            node = anchor[0]
        if node is not None:
            extent = node.find(f'.//{wp_ns}extent')
            if extent is not None:
                try:
                    cx = int(extent.get('cx')); cy = int(extent.get('cy'))
                except Exception:
                    cx = cy = None

        results.append((blob, cx, cy, content_type))
    return results

def _maybe_convert_with_pillow(blob: bytes, _mime: str) -> Optional[bytes]:
    if Image is None:
        return None
    try:
        with Image.open(BytesIO(blob)) as im:
            out = BytesIO()
            im.save(out, format="PNG")
            return out.getvalue()
    except Exception:
        return None

def _compute_target_width_emu(cx: Optional[int], blob: bytes) -> Emu:
    if cx:
        width_cm = cx / 360000.0
        return Emu(Cm(SMALL_WIDTH_CM if width_cm < SMALL_SOURCE_THRESHOLD_CM else NORMAL_WIDTH_CM))
    if Image is not None:
        try:
            with Image.open(BytesIO(blob)) as im:
                px_w = im.width
                dpi = im.info.get("dpi", (96, 96))[0] or 96
                width_cm = (px_w / dpi) * 2.54
                return Emu(Cm(SMALL_WIDTH_CM if width_cm < SMALL_SOURCE_THRESHOLD_CM else NORMAL_WIDTH_CM))
        except Exception:
            pass
    return Emu(Cm(NORMAL_WIDTH_CM))

def _apply_border_and_shadow_to_pic(pic_el) -> None:
    spPr_list = pic_el.xpath('.//pic:spPr')
    if spPr_list:
        spPr = spPr_list[0]
    else:
        spPr = parse_xml(f'<pic:spPr {nsdecls("pic")} {nsdecls("a")}></pic:spPr>')
        pic_el.insert(0, spPr)

    if not spPr.xpath('.//a:prstGeom'):
        prstGeom = parse_xml(f'<a:prstGeom prst="rect" {nsdecls("a")}><a:avLst/></a:prstGeom>')
        spPr.append(prstGeom)

    for ln in spPr.xpath('.//a:ln'):
        parent = ln.getparent()
        if parent is not None:
            parent.remove(ln)
    for eff in spPr.xpath('.//a:effectLst'):
        parent = eff.getparent()
        if parent is not None:
            parent.remove(eff)

    ln_xml = f'''
    <a:ln w="19050" cap="flat" cmpd="sng" algn="ctr" {nsdecls('a')}>
      <a:solidFill><a:srgbClr val="404040"/></a:solidFill>
      <a:miter lim="800000"/>
      <a:headEnd type="none" w="med" len="med"/>
      <a:tailEnd type="none" w="med" len="med"/>
    </a:ln>
    '''
    spPr.append(parse_xml(ln_xml))

    shdw_xml = f'''
    <a:effectLst {nsdecls('a')}>
      <a:outerShdw blurRad="6350" dist="3810" dir="2700000" algn="b" rotWithShape="0">
        <a:srgbClr val="404040"><a:alpha val="40000"/></a:srgbClr>
      </a:outerShdw>
    </a:effectLst>
    '''
    spPr.append(parse_xml(shdw_xml))

def _insert_image_block_with_caption(container: Union[Document, _Cell],
                                     blob: bytes,
                                     cx: Optional[int],
                                     figure_num: int) -> None:
    width_emu = _compute_target_width_emu(cx, blob)

    img_p = container.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = img_p.add_run()
    try:
        run_img.add_picture(BytesIO(blob), width=width_emu)
    except UnrecognizedImageError:
        return
    except Exception:
        return

    try:
        pics = run_img._r.xpath('.//pic:pic')
        if pics:
            _apply_border_and_shadow_to_pic(pics[-1])
    except Exception:
        pass

    cap_p = container.add_paragraph(f"{CAPTION_PREFIX} {figure_num}")
    try:
        cap_p.style = container.part.document.styles.get(CAPTION_STYLE_NAME, cap_p.style)
    except Exception:
        pass
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _insert_images_after_paragraph(container: Union[Document, _Cell],
                                   images: List[tuple],
                                   figure_counter: List[int]) -> Tuple[int, int]:
    inserted = skipped = 0
    for blob, cx, _cy, mime in images:
        use_blob = blob
        if (mime or "").lower() not in SUPPORTED_IMG_MIMES:
            converted = _maybe_convert_with_pillow(blob, mime or "")
            if not converted:
                skipped += 1
                continue
            use_blob = converted

        figure_counter[0] += 1
        try:
            _insert_image_block_with_caption(container, use_blob, cx, figure_counter[0])
            inserted += 1
        except Exception:
            skipped += 1
    return inserted, skipped

# ===================== COPY HELPERS =====================

def _copy_runs(dst_para, src_para) -> None:
    if src_para.runs:
        for r in src_para.runs:
            nr = dst_para.add_run(r.text)
            nr.bold = r.bold
            nr.italic = r.italic
            nr.underline = r.underline
    else:
        dst_para.add_run("")

def _copy_paragraph(dst_doc: Document, src_para, style_name: str,
                    raw_doc: Document, figure_counter: List[int]) -> Tuple[int, int]:
    p = dst_doc.add_paragraph()
    try:
        p.style = dst_doc.styles[style_name]
    except Exception:
        p.style = style_name
    _copy_runs(p, src_para)

    inserted = skipped = 0
    for run in src_para.runs:
        images = _extract_inline_images_from_run(raw_doc, run)
        if images:
            i, s = _insert_images_after_paragraph(dst_doc, images, figure_counter)
            inserted += i
            skipped += s
    return inserted, skipped

def _copy_cell(dst_cell: _Cell, src_cell: _Cell, tpl: Document,
               raw_doc: Document, figure_counter: List[int]) -> Tuple[int, int]:
    for p in list(dst_cell.paragraphs):
        p._element.getparent().remove(p._element)

    inserted_total = skipped_total = 0

    for p in src_cell.paragraphs:
        style_name = _choose_style_for_paragraph(tpl, p)
        new_p = dst_cell.add_paragraph()
        try:
            new_p.style = tpl.styles[style_name]
        except Exception:
            new_p.style = style_name
        _copy_runs(new_p, p)

        for run in p.runs:
            images = _extract_inline_images_from_run(raw_doc, run)
            if not images:
                continue
            i, s = _insert_images_after_paragraph(dst_cell, images, figure_counter)
            inserted_total += i
            skipped_total += s

    return inserted_total, skipped_total

def _copy_table(dst_doc: Document, src_table: Table, tpl: Document,
                raw_doc: Document, figure_counter: List[int]) -> Tuple[int, int]:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    if rows == 0 or cols == 0:
        return 0, 0

    new_tbl = dst_doc.add_table(rows=rows, cols=cols, style=TABLE_STYLE_NAME or None)

    inserted_total = skipped_total = 0
    for r in range(rows):
        for c in range(cols):
            i, s = _copy_cell(new_tbl.cell(r, c), src_table.cell(r, c), tpl, raw_doc, figure_counter)
            inserted_total += i
            skipped_total += s
    return inserted_total, skipped_total

# ===================== PUBLIC API =====================

def merge_into_template(template_path: str, raw_docx_path: str, out_path: str) -> Dict[str, int]:
    """
    Merge raw_docx into template and save to out_path.
    Returns a dict with inserted/skipped image counts:
      { "inserted_images": int, "skipped_images": int }
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not os.path.exists(raw_docx_path):
        raise FileNotFoundError(f"Raw DOCX not found: {raw_docx_path}")

    tpl = Document(template_path)
    raw = Document(raw_docx_path)

    _append_page_break(tpl)

    figure_counter = [0]
    skipped_images = 0
    inserted_total = 0

    for element in raw.element.body:
        if element.tag == qn('w:p'):
            para = next((p for p in raw.paragraphs if p._p is element), None)
            if para is None:
                continue
            try:
                if para.paragraph_format.page_break_before:
                    _append_page_break(tpl)
            except Exception:
                pass

            style_name = _choose_style_for_paragraph(tpl, para)
            i, skipped = _copy_paragraph(tpl, para, style_name, raw, figure_counter)
            inserted_total += i
            skipped_images += skipped

        elif element.tag == qn('w:tbl'):
            src_table = next((t for t in raw.tables if t._tbl is element), None)
            if src_table is not None:
                i, skipped = _copy_table(tpl, src_table, tpl, raw, figure_counter)
                inserted_total += i
                skipped_images += skipped

    _set_update_fields_on_open(tpl)

    if skipped_images:
        note = tpl.add_paragraph(f"[note] {skipped_images} unsupported image(s) were skipped.")
        try:
            note.style = tpl.styles.get(DEFAULT_BODY_STYLE, note.style)
        except Exception:
            pass

    tpl.save(out_path)

    return {"inserted_images": inserted_total, "skipped_images": skipped_images}


def merge_from_any(template_path: str, raw_path: str, out_docx: str) -> Dict[str, int]:
    """
    Accepts .docx or .md
      - .md  → temporary .docx via parser.md_file_to_docx(), then merge
      - .docx → merge directly

    Returns a dict with image stats: {"inserted_images": int, "skipped_images": int}
    """
    src = raw_path
    ext = Path(raw_path).suffix.lower()
    if ext in {".md", ".markdown", ".mdx"}:
        tmp_docx = str(Path(out_docx).with_name(Path(out_docx).stem + "_from_md.docx"))
        md_file_to_docx(raw_path, tmp_docx)
        src = tmp_docx

    stats = merge_into_template(template_path, src, out_docx)
    return stats
