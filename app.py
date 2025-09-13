# app.py
import time
from pathlib import Path
from datetime import datetime

from flask import Flask, request, send_from_directory, send_file, url_for, render_template, abort, make_response
from werkzeug.utils import secure_filename
import json

from merge import merge_from_any
from converters import detect_pdf_engine, docx_to_pdf
from parser import md_file_to_docx   # unchanged

# for reading core properties and counting PDF pages
from docx import Document as DocxDocument
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

app = Flask(__name__)
app.secret_key = "dev-secret"

BASE_DIR   = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
for p in (UPLOAD_DIR, OUTPUT_DIR):
    p.mkdir(exist_ok=True)

# ~100 MB request cap
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024

ALLOWED_RAW = {".docx", ".md", ".markdown", ".mdx"}
ALLOWED_MD  = {".md", ".markdown", ".mdx"}
ALLOWED_TPL = {".docx"}


def _ext_ok(filename: str, allowed: set[str]) -> bool:
    return "." in filename and Path(filename).suffix.lower() in allowed


def sizeof_fmt(num: int) -> str:
    """Human-readable file size."""
    for unit in ['B','KB','MB','GB','TB']:
        if num < 1024.0:
            return f"{num:3.1f} {unit}"
        num /= 1024.0
    return f"{num:.1f} PB"


def _icon_for_name(name: str) -> str:
    ext = Path(name).suffix.lower()
    if ext in ('.md', '.markdown', '.mdx'):
        return 'ti ti-markdown'
    if ext in ('.doc', '.docx'):
        return 'ti ti-file-type-docx'
    if ext == '.pdf':
        return 'ti ti-file-type-pdf'
    if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.svg'):
        return 'ti ti-image'
    return 'ti ti-file-text'


@app.get("/")
def index():
    return render_template(
        "index.html",
        docx_url="",
        pdf_url="",
        pdf_preview_url="",
        pdf_error_message="",
        doc_version="",
        issued_date="",
        doc_author="",
        template_name="",
        total_pages="",
        description="",
        docx_size="",
        pdf_size="",
        conversion_time="",
        image_count="",
        skipped_images="",
        uploaded_files=[],
        server_message=""
    )


@app.post("/convert")
def convert():
    tpl        = request.files.get("template_file")
    raw_single = request.files.get("raw_file")
    raw_many   = request.files.getlist("raw_files") or []

    if not tpl or not tpl.filename or not _ext_ok(tpl.filename, ALLOWED_TPL):
        return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message="Missing or invalid template", doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=[], server_message="Missing or invalid template")

    if not raw_single and not raw_many:
        return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message="No Markdown/Raw file(s) uploaded", doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=[], server_message="No Markdown/Raw file(s) uploaded")

    if raw_many and raw_single:
        raw_single = None

    if raw_many:
        filtered = []
        for f in raw_many:
            if not f or not f.filename:
                continue
            if not _ext_ok(f.filename, ALLOWED_MD):
                return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message="Invalid raw file in list", doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=[], server_message="Invalid raw file in list")
            filtered.append(f)
        if not filtered:
            return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message="No valid markdown files", doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=[], server_message="No valid markdown files")
        if len(filtered) > 20:
            filtered = filtered[:20]
        raw_many = filtered

    if raw_single:
        if not raw_single.filename or not _ext_ok(raw_single.filename, ALLOWED_RAW):
            return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message="Invalid raw file", doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=[], server_message="Invalid raw file")

    ts = time.strftime("%Y%m%d-%H%M%S")

    safe_tpl  = secure_filename(tpl.filename) or f"template_{ts}.docx"
    tpl_path  = UPLOAD_DIR / f"{Path(safe_tpl).stem}_{ts}.docx"
    tpl.save(tpl_path)

    out_docx  = OUTPUT_DIR / f"merged_{ts}.docx"
    out_pdf   = OUTPUT_DIR / f"merged_{ts}.pdf"

    apply_img_style = request.form.get("img_style") is not None

    pdf_error_message = ""
    image_count = 0
    skipped_images = 0
    conversion_time = ""
    uploaded_files = []

    try:
        t0 = time.perf_counter()

        if raw_many:
            saved = []
            base_dirs = []
            saved_pairs = []
            for i, f in enumerate(raw_many, start=1):
                original_name = f.filename or f"section_{i:02d}.md"
                safe = secure_filename(original_name) or f"section_{i:02d}.md"
                p = UPLOAD_DIR / f"{Path(safe).stem}_{ts}_{i:02d}{Path(safe).suffix.lower()}"
                f.save(p)
                saved.append(p)
                saved_pairs.append((p, original_name))
                base_dirs.append(Path(original_name).parent)

            for p, orig in saved_pairs:
                try:
                    size = p.stat().st_size
                except Exception:
                    size = 0
                uploaded_files.append({
                    "name": str(p.name),
                    "display_name": orig,
                    "size": size,
                    "size_human": sizeof_fmt(size),
                    "icon": _icon_for_name(orig)
                })

            combined_md = UPLOAD_DIR / f"combined_{ts}.md"
            with combined_md.open("w", encoding="utf-8") as out:
                for i, p in enumerate(saved, start=1):
                    text = p.read_text(encoding="utf-8")
                    out.write(text)
                    if i != len(saved):
                        out.write("\n\n")

            tmp_docx = str(out_docx.with_name(out_docx.stem + "_from_md.docx"))
            md_file_to_docx(
                str(combined_md),
                tmp_docx,
                base_dir=base_dirs[0].resolve() if base_dirs else None,
                style_images=apply_img_style,
                debug=True,
            )

            stats = merge_from_any(str(tpl_path), tmp_docx, str(out_docx))
        else:
            raw_ext  = Path(raw_single.filename).suffix.lower()
            original_name = raw_single.filename or f"raw_{ts}{raw_ext}"
            safe_raw = secure_filename(original_name) or f"raw_{ts}{raw_ext}"
            raw_path = UPLOAD_DIR / f"{Path(safe_raw).stem}_{ts}{raw_ext}"
            raw_single.save(raw_path)

            try:
                size = raw_path.stat().st_size
            except Exception:
                size = 0
            uploaded_files.append({
                "name": str(raw_path.name),
                "display_name": original_name,
                "size": size,
                "size_human": sizeof_fmt(size),
                "icon": _icon_for_name(original_name)
            })

            if raw_ext in {".md", ".markdown", ".mdx"}:
                md_file_to_docx(
                    str(raw_path),
                    str(out_docx),
                    base_dir=Path(original_name).parent,
                    style_images=apply_img_style,
                    debug=True,
                )
                stats = merge_from_any(str(tpl_path), str(out_docx), str(out_docx))
            else:
                stats = merge_from_any(str(tpl_path), str(raw_path), str(out_docx))

        t1 = time.perf_counter()
        conversion_time = f"{(t1 - t0):.2f}s"

        image_count = int(stats.get("inserted_images", 0))
        skipped_images = int(stats.get("skipped_images", 0))

    except Exception as e:
        print("Merge error:", e)
        pdf_error_message = "Internal error while merging files."
        return render_template("index.html", docx_url="", pdf_url="", pdf_preview_url="", pdf_error_message=pdf_error_message, doc_version="", issued_date="", doc_author="", template_name="", total_pages="", description="", docx_size="", pdf_size="", conversion_time="", image_count="", skipped_images="", uploaded_files=uploaded_files, server_message="Internal error while merging files.")

    docx_url = url_for("download_output", filename=out_docx.name) if out_docx.exists() else ""

    pdf_url = ""
    pdf_preview_url = ""
    available, detail = detect_pdf_engine()
    if available:
        err = docx_to_pdf(out_docx, out_pdf)
        if err is None and out_pdf.exists():
            pdf_url = url_for("download_output", filename=out_pdf.name)
            pdf_preview_url = url_for("preview_output", filename=out_pdf.name)
        else:
            print("PDF conversion failed:", err)
            pdf_error_message = err or "PDF conversion failed on the server."
    else:
        print("No PDF engine detected:", detail)
        pdf_error_message = "No PDF converter installed (LibreOffice or MS Word required)."

    doc_version = ""
    issued_date = ""
    doc_author = ""
    template_name = tpl.filename if tpl and tpl.filename else ""
    total_pages = ""
    description = ""

    try:
        doc_tpl = DocxDocument(str(tpl_path))
        cp = doc_tpl.core_properties

        found = False
        try:
            for tbl in doc_tpl.tables:
                if len(tbl.rows) < 2:
                    continue
                header_cells = [c.text.strip() for c in tbl.rows[0].cells]
                header_norm = [h.lower().strip() for h in header_cells]
                if any('version' in h for h in header_norm) and (any('author' in h for h in header_norm) or any('issued' in h for h in header_norm) or any('description' in h for h in header_norm)):
                    data_cells = [c.text.strip() for c in tbl.rows[1].cells]
                    hdr_to_val = {}
                    for idx, h in enumerate(header_norm):
                        val = data_cells[idx] if idx < len(data_cells) else ""
                        hdr_to_val[h] = val

                    def _get_by_variants(variants):
                        for v in variants:
                            for k, val in hdr_to_val.items():
                                if v in k:
                                    return val
                        return None

                    description = _get_by_variants(['description', 'desc']) or description
                    doc_version = _get_by_variants(['version', 'rev']) or doc_version
                    issued_date_val = _get_by_variants(['issued date', 'issued', 'issue date', 'date'])
                    if issued_date_val:
                        issued_date = issued_date_val
                    doc_author = _get_by_variants(['author', 'creator']) or doc_author
                    total_pages = _get_by_variants(['total pages', 'pages']) or total_pages
                    found = True
                    break
        except Exception as e:
            print("Revision-table parsing error:", e)

        if not found:
            found_map = {}
            try:
                for tbl in doc_tpl.tables:
                    for row in tbl.rows:
                        if len(row.cells) < 2:
                            continue
                        left = row.cells[0].text.strip()
                        right = row.cells[1].text.strip()
                        if not left:
                            continue
                        key = left.lower().strip().rstrip(':')
                        if key:
                            found_map[key] = right
                    if any(k in found_map for k in ("version", "issued date", "issued", "author", "template", "total pages", "description")):
                        break
            except Exception as e:
                print("Table parse error (fallback):", e)

            if found_map:
                description = found_map.get("description") or found_map.get("desc") or description
                doc_version = found_map.get("version") or found_map.get("rev") or doc_version
                issued_date = found_map.get("issued date") or found_map.get("issued") or found_map.get("issued_date") or issued_date
                doc_author = found_map.get("author") or found_map.get("creator") or doc_author
                template_name = found_map.get("template") or template_name
                total_pages = found_map.get("total pages") or found_map.get("pages") or total_pages

        if not doc_author:
            author = getattr(cp, "author", None) or getattr(cp, "creator", None) or ""
            doc_author = str(author) if author else doc_author
        if not issued_date:
            created = getattr(cp, "created", None)
            if created:
                if isinstance(created, str):
                    issued_date = created
                else:
                    try:
                        issued_date = created.strftime("%Y-%m-%d %H:%M:%S")
                    except Exception:
                        issued_date = str(created)
        if not doc_version:
            revision = getattr(cp, "revision", None)
            if revision:
                doc_version = str(revision)
    except Exception as e:
        print("Template metadata read failed:", e)

    try:
        if pdf_url and PdfReader is not None and out_pdf.exists():
            try:
                reader = PdfReader(str(out_pdf))
                total_pages = str(len(reader.pages))
            except Exception as e:
                print("PDF page count failed:", e)
    except Exception as e:
        print("page count/check error:", e)

    docx_size = ""
    pdf_size = ""
    try:
        if out_docx.exists():
            docx_size = sizeof_fmt(out_docx.stat().st_size)
    except Exception:
        docx_size = ""
    try:
        if out_pdf.exists():
            pdf_size = sizeof_fmt(out_pdf.stat().st_size)
    except Exception:
        pdf_size = ""

    return render_template(
        "index.html",
        docx_url=docx_url,
        pdf_url=pdf_url,
        pdf_preview_url=pdf_preview_url,
        pdf_error_message=pdf_error_message,
        doc_version=doc_version,
        issued_date=issued_date,
        doc_author=doc_author,
        template_name=template_name,
        total_pages=total_pages,
        description=description,
        docx_size=docx_size,
        pdf_size=pdf_size,
        conversion_time=conversion_time,
        image_count=str(image_count),
        skipped_images=str(skipped_images),
        uploaded_files=uploaded_files,
        server_message=""
    )


@app.get("/outputs/<path:filename>")
def download_output(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


@app.get("/outputs/inline/<path:filename>")
def preview_output(filename):
    target = OUTPUT_DIR / filename
    if not target.exists():
        abort(404)

    if target.suffix.lower() == ".pdf":
        try:
            resp = make_response(send_file(str(target.resolve()), mimetype="application/pdf"))
            resp.headers["Content-Disposition"] = f"inline; filename={target.name}"
            resp.headers['X-Frame-Options'] = 'SAMEORIGIN'
            resp.headers['Content-Security-Policy'] = "frame-ancestors 'self'"
            return resp
        except Exception as e:
            print("Preview serve error:", e)
            return send_from_directory(OUTPUT_DIR, filename, as_attachment=False)

    resp = make_response(send_from_directory(OUTPUT_DIR, filename, as_attachment=False))
    resp.headers['X-Frame-Options'] = 'SAMEORIGIN'
    resp.headers['Content-Security-Policy'] = "frame-ancestors 'self'"
    return resp


if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=5055)
