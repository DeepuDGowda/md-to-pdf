"""
Markdown → python-docx:
  • #, ##, ###, ####  → Word Heading 1..4
  • **bold**          → real bold runs
  • - item            → List Bullet / List Bullet 2
  • Pipe tables with header separator → Word tables
  • ![alt](path)      → Embedded image with caption (if meaningful)

This version aggressively strips:
 - YAML front matter at the top of a file (--- ... ---)
 - custom hidden blocks starting with '--- text ---' up to a following
   '<!-- ### Key Activities -->' marker
 - ALL HTML comments <!-- ... --> (single-line or multi-line)
before parsing so they do not appear in the output DOCX/PDF.
"""

from __future__ import annotations
import re, json
from pathlib import Path
from typing import List
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Emu
from docx.table import _Cell
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

try:
    from PIL import Image
except Exception:
    Image = None


# -------- Load config.json --------
CONFIG_FILE = Path(__file__).parent / "config.json"
GLOBAL_IMAGE_DIRS: List[Path] = []

if CONFIG_FILE.exists():
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            dirs = data.get("GLOBAL_IMAGE_DIRS", [])
            GLOBAL_IMAGE_DIRS = [Path(d).resolve() for d in dirs if d.strip()]
    except Exception as e:
        print("⚠️ Failed to read config.json:", e)

# -------- inline **bold** --------
BOLD = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)

def _emit_bold_runs(p, text: str) -> None:
    if not text:
        p.add_run("")
        return
    i = 0
    for m in BOLD.finditer(text):
        if m.start() > i:
            p.add_run(text[i:m.start()])
        r = p.add_run(m.group(1))
        r.bold = True
        i = m.end()
    if i < len(text):
        p.add_run(text[i:])

# -------- Hidden / comment stripping --------
# 1) YAML front matter at the top: --- ... ---
_YAML_FRONT_MATTER_RE = re.compile(r'(?s)^\s*---\s*\n.*?\n---\s*\n', flags=re.IGNORECASE)

# 2) Custom hidden block: starts with a line like '--- text ---' and ends at
#    the following HTML comment that contains 'Key Activities' (e.g. <!-- ### Key Activities -->).
_HIDDEN_SECTION_RE = re.compile(
    r'(?mis)^[ \t]*---[ \t]*text[ \t]*---[ \t\r\n]*.*?<!--\s*#{0,}\s*Key\s+Activities\s*-->[ \t\r\n]*',
    flags=re.IGNORECASE
)

# 3) Remove all HTML comments anywhere
_HTML_COMMENT_RE = re.compile(r'(?s)<!--.*?-->', flags=re.IGNORECASE)

def _strip_hidden_sections(md_text: str) -> str:
    """
    Remove YAML front matter, custom hidden sections and all HTML comments.
    Returns cleaned markdown text.
    """
    if not md_text:
        return md_text

    cleaned = md_text

    # Remove YAML front matter at the top (if present)
    cleaned, n_subs_yaml = _YAML_FRONT_MATTER_RE.subn("", cleaned)

    # Remove custom hidden blocks like '--- text ---' ... '<!-- ### Key Activities -->'
    cleaned, n_subs_hidden = _HIDDEN_SECTION_RE.subn("", cleaned)

    # Remove ALL HTML comments (single/multi-line) - forceful removal
    cleaned, n_subs_comments = _HTML_COMMENT_RE.subn("", cleaned)

    # Optionally, trim leading/trailing whitespace/newlines
    cleaned = cleaned.lstrip("\n\r ").rstrip()

    # Debug prints (comment out if noisy)
    if any((n_subs_yaml, n_subs_hidden, n_subs_comments)):
        print(f"🔒 Hidden removal: yaml={n_subs_yaml}, custom_hidden={n_subs_hidden}, html_comments={n_subs_comments}")

    return cleaned

# -------- pipe tables --------
_CELL_SEP = re.compile(r"^\s*:?-{3,}:?\s*$")

def _is_table_separator_row(line: str) -> bool:
    if "|" not in line:
        return False
    tokens = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(_CELL_SEP.match(tok) is not None for tok in tokens)

def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]

def _align_cell(cell: _Cell, align: str) -> None:
    par = cell.paragraphs[0]
    par.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

def _parse_alignments(sep_line: str) -> List[str]:
    aligns: List[str] = []
    for token in _split_row(sep_line):
        left  = token.startswith(":")
        right = token.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns

def _render_table(doc: Document, header_line: str, sep_line: str, body_lines: List[str]) -> None:
    hdr = _split_row(header_line)
    aligns = _parse_alignments(sep_line)
    rows   = [_split_row(r) for r in body_lines]

    n_cols = max([len(hdr)] + [len(r) for r in rows]) if rows else len(hdr)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)

    # header row
    for j in range(n_cols):
        cell = tbl.cell(0, j)
        txt = hdr[j] if j < len(hdr) else ""
        _emit_bold_runs(cell.paragraphs[0], txt)
        for run in cell.paragraphs[0].runs:
            run.bold = True
        if j < len(aligns):
            _align_cell(cell, aligns[j])

    # body rows
    for i, row in enumerate(rows, start=1):
        for j in range(n_cols):
            cell = tbl.cell(i, j)
            txt = row[j] if j < len(row) else ""
            _emit_bold_runs(cell.paragraphs[0], txt)
            if j < len(aligns):
                _align_cell(cell, aligns[j])

# -------- Images --------
IMG = re.compile(r"!\[(.*?)\]\((.*?)\)")

CAPTION_PREFIX = "Image"
CAPTION_STYLE_NAME = "Image Description"
NORMAL_WIDTH_CM = 15.0
SMALL_WIDTH_CM  = 9.0
SMALL_SOURCE_THRESHOLD_CM = 10.0

ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".svg", ".gif", ".ico", ".bmp", ".tiff"}

def _compute_width(path: Path) -> Emu:
    if Image is not None and path.exists() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}:
        try:
            with Image.open(path) as im:
                px_w = im.width
                dpi = im.info.get("dpi", (96, 96))[0] or 96
                width_cm = (px_w / dpi) * 2.54
                return Emu(Cm(SMALL_WIDTH_CM if width_cm < SMALL_SOURCE_THRESHOLD_CM else NORMAL_WIDTH_CM))
        except Exception:
            pass
    return Emu(Cm(NORMAL_WIDTH_CM))

def _apply_border_and_shadow_to_pic(pic_el) -> None:
    spPr_list = pic_el.xpath('.//pic:spPr')
    if spPr_list:
        spPr = spPr_list[0]
    else:
        spPr = parse_xml(f'<pic:spPr {nsdecls("pic")} {nsdecls("a")}></pic:spPr>')
        pic_el.insert(0, spPr)

    # Border
    ln_xml = f'''
    <a:ln w="19050" cap="flat" cmpd="sng" algn="ctr" {nsdecls('a')}>
      <a:solidFill><a:srgbClr val="404040"/></a:solidFill>
      <a:miter lim="800000"/>
    </a:ln>
    '''
    spPr.append(parse_xml(ln_xml))

    # Shadow
    shdw_xml = f'''
    <a:effectLst {nsdecls('a')}>
      <a:outerShdw blurRad="6350" dist="3810" dir="2700000" algn="b" rotWithShape="0">
        <a:srgbClr val="404040"><a:alpha val="40000"/></a:srgbClr>
      </a:outerShdw>
    </a:effectLst>
    '''
    spPr.append(parse_xml(shdw_xml))


# -------- Path Normalization + Config + Recursive --------
def _resolve_image_path(path_str: str, base_dir: Path | None) -> Path:
    """
    Resolve image path with guaranteed recursive search in GLOBAL_IMAGE_DIRS.
    """
    path_str = path_str.strip().replace("\\", "/")
    while path_str.startswith("/"):
        path_str = path_str.lstrip("/")

    img_path = Path(path_str)
    print(f"\n🔍 Resolving image: {path_str}")

    candidates = []

    # --- Step 1: Relative to Markdown file ---
    if base_dir:
        candidates.append((base_dir / img_path).resolve())
        candidates.append((base_dir / "images" / img_path.name).resolve())
        candidates.append((base_dir / "img" / img_path.name).resolve())

    # --- Step 2: Project root ---
    project_root = Path.cwd()
    candidates.append((project_root / img_path).resolve())
    candidates.append((project_root / "images" / img_path.name).resolve())
    candidates.append((project_root / "static" / "img" / img_path.name).resolve())

    # --- Step 3: User-configured global dirs (direct paths) ---
    for gdir in GLOBAL_IMAGE_DIRS:
        candidates.append((gdir / img_path).resolve())
        candidates.append((gdir / "images" / img_path.name).resolve())
        candidates.append((gdir / "static" / "img" / img_path.name).resolve())

    # Check candidates first
    for c in candidates:
        print("   checking:", c, "✅" if c.exists() else "❌")
    for c in candidates:
        if c.exists():
            print("👉 Found:", c)
            return c

    # --- Step 4: Always do global recursive search ---
    for gdir in GLOBAL_IMAGE_DIRS:
        if gdir.exists():
            for p in gdir.rglob("*"):
                if p.suffix.lower() in ALLOWED_EXTENSIONS and p.name.lower() == img_path.name.lower():
                    print("🔎 Recursive match (global):", p)
                    return p

    # Nothing found
    print("⚠️ Not found, returning raw path:", img_path.resolve())
    return img_path


def _insert_image(doc: Document, path: Path, alt: str, style_images: bool) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Image not found: {path.resolve()}]")
        return

    width = _compute_width(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    try:
        run.add_picture(str(path), width=width)
        if style_images:
            pics = run._r.xpath('.//pic:pic')
            if pics:
                _apply_border_and_shadow_to_pic(pics[-1])
    except Exception:
        doc.add_paragraph(f"[Unsupported image: {path.resolve()}]")
        return

    # 🔥 Only add caption if alt is meaningful
    if alt and alt.strip().lower() != "image":
        caption_text = f"{CAPTION_PREFIX}: {alt}".strip()
        cap = doc.add_paragraph(caption_text)
        try:
            cap.style = doc.styles[CAPTION_STYLE_NAME]
        except Exception:
            pass
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# -------- block parsing --------
H_HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
B_BULLET  = re.compile(r"^(?P<indent>\s*)-\s+(?P<text>.*\S)\s*$")

def render_markdown_into(
    doc: Document, md_text: str, base_dir: Path | None = None,
    style_images: bool = True
) -> Document:
    # Preprocess: strip hidden/internal sections, YAML front matter and HTML comments
    md_text = _strip_hidden_sections(md_text)

    lines = md_text.splitlines()
    i = 0
    para_buf: List[str] = []

    def flush_para():
        nonlocal para_buf
        if not para_buf:
            return
        text = " ".join(s.strip() for s in para_buf).strip()
        p = doc.add_paragraph()
        _emit_bold_runs(p, text)
        para_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip() == "":
            flush_para()
            doc.add_paragraph("")
            i += 1
            continue

        # heading
        m = H_HEADING.match(line)
        if m:
            flush_para()
            hashes, txt = m.groups()
            level = min(len(hashes), 4)
            p = doc.add_paragraph()
            try:
                p.style = f"Heading {level}"
            except Exception:
                pass
            _emit_bold_runs(p, txt)
            i += 1
            continue

        # bullet
        m = B_BULLET.match(line)
        if m:
            flush_para()
            while i < len(lines):
                bm = B_BULLET.match(lines[i])
                if not bm:
                    break
                indent = len(bm.group("indent") or "")
                text   = bm.group("text")
                bp = doc.add_paragraph()
                try:
                    bp.style = "List Bullet 2" if indent >= 1 else "List Bullet"
                except Exception:
                    pass
                _emit_bold_runs(bp, text)
                i += 1
            continue

        # table
        if "|" in line and i + 1 < len(lines) and _is_table_separator_row(lines[i + 1]):
            flush_para()
            header = line
            sep = lines[i + 1]
            i += 2
            body: List[str] = []
            while i < len(lines):
                l = lines[i]
                if l.strip() == "" or "|" not in l:
                    break
                body.append(l)
                i += 1
            _render_table(doc, header, sep, body)
            continue

        # image
        m = IMG.search(line)
        if m:
            flush_para()
            alt, path_str = m.groups()
            img_path = _resolve_image_path(path_str, base_dir)
            _insert_image(doc, img_path, alt, style_images)
            i += 1
            continue

        # normal text
        para_buf.append(line)
        i += 1

    flush_para()
    return doc


def md_file_to_docx(
    md_path: str,
    out_docx: str,
    base_dir: Path | None = None,
    style_images: bool = True,
    debug: bool = False
) -> None:
    md_file = Path(md_path)
    text = md_file.read_text(encoding="utf-8")
    # Defensive: strip hidden sections here as well
    text = _strip_hidden_sections(text)
    doc = Document()
    render_markdown_into(
        doc,
        text,
        base_dir or md_file.parent,
        style_images=style_images
    )
    doc.save(out_docx)
    if debug:
        print(f"✅ Saved: {out_docx}")
